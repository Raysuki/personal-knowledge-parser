import base64
import hashlib
import io
import json
import logging
import os
import re
import sqlite3
import tempfile
import urllib.parse
import zipfile
from datetime import datetime
from typing import Any, Dict, List, Optional, Tuple

import fitz
import httpx
import numpy as np
from docx import Document
from dotenv import load_dotenv
from fastapi import FastAPI, File, Form, HTTPException, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse, StreamingResponse
from PIL import Image
from pydantic import BaseModel
from volcenginesdkarkruntime import Ark


load_dotenv()

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

PROJECT_DIR = os.path.dirname(os.path.abspath(__file__))
DB_PATH = os.path.join(PROJECT_DIR, "knowledge_base.db")
os.environ.setdefault("PADDLE_PDX_CACHE_HOME", os.path.join(PROJECT_DIR, ".paddlex-cache"))
os.environ.setdefault("PADDLE_PDX_DISABLE_MODEL_SOURCE_CHECK", "True")

app = FastAPI(title="Experience Proof Extract API")

app.add_middleware(
    CORSMiddleware,
    allow_origins=[
        "http://localhost:3000",
        "http://127.0.0.1:3000",
    ],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

ARK_API_KEY = os.getenv("ARK_API_KEY")
ARK_BASE_URL = os.getenv("ARK_BASE_URL", "https://ark.cn-beijing.volces.com/api/v3")
ARK_TEXT_ENDPOINT_ID = (
    os.getenv("ARK_TEXT_ENDPOINT_ID")
    or os.getenv("DOUBAO_TEXT_MODEL")
    or os.getenv("ARK_ENDPOINT_ID")
    or os.getenv("DOUBAO_VISION_MODEL")
)
ARK_VISION_ENDPOINT_ID = (
    os.getenv("ARK_VISION_ENDPOINT_ID")
    or os.getenv("DOUBAO_VISION_MODEL")
    or os.getenv("ARK_ENDPOINT_ID")
)
ARK_TIMEOUT = int(os.getenv("ARK_TIMEOUT", "180"))
ARK_MAX_TOKENS = int(os.getenv("ARK_MAX_TOKENS", "4000"))

OCR_LANG = os.getenv("OCR_LANG", "ch")
PDF_TEXT_THRESHOLD = int(os.getenv("PDF_TEXT_THRESHOLD", "80"))
MAX_TEXT_CHARS = int(os.getenv("MAX_TEXT_CHARS", "12000"))

SUPPORTED_IMAGE_TYPES = {"image/png", "image/jpeg"}
SUPPORTED_PDF_TYPES = {"application/pdf"}
SUPPORTED_DOCX_TYPES = {
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
}
SUPPORTED_TEXT_TYPES = {"text/plain"}

_ocr_engine = None


class KnowledgeSavePayload(BaseModel):
    knowledge_base_id: int
    doc_id: Optional[int] = None
    user_id: Optional[str] = None
    file_name: Optional[str] = None
    file_hash: Optional[str] = None
    basic_info: Dict[str, Any]
    education_history: List[Dict[str, Any]]
    experiences: List[Dict[str, Any]]


class KnowledgeBaseCreatePayload(BaseModel):
    name: str


class KnowledgeBaseRenamePayload(BaseModel):
    name: str


@app.on_event("startup")
async def startup_event() -> None:
    init_db()
    logger.info("ARK_BASE_URL=%s", ARK_BASE_URL)
    logger.info("ARK_TEXT_ENDPOINT_ID=%s", ARK_TEXT_ENDPOINT_ID)
    logger.info("ARK_VISION_ENDPOINT_ID=%s", ARK_VISION_ENDPOINT_ID)
    logger.info("ARK_API_KEY loaded=%s", bool(ARK_API_KEY))
    logger.info("ARK_TIMEOUT=%s", ARK_TIMEOUT)
    logger.info("ARK_MAX_TOKENS=%s", ARK_MAX_TOKENS)
    logger.info("OCR_LANG=%s", OCR_LANG)


def get_db_connection() -> sqlite3.Connection:
    connection = sqlite3.connect(DB_PATH)
    connection.row_factory = sqlite3.Row
    return connection


def init_db() -> None:
    connection = get_db_connection()
    try:
        connection.execute(
            """
            CREATE TABLE IF NOT EXISTS knowledge_bases (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                user_id TEXT,
                name TEXT NOT NULL UNIQUE,
                basic_info TEXT,
                education_history TEXT,
                experiences TEXT,
                created_at TEXT NOT NULL,
                updated_at TEXT NOT NULL
            )
            """
        )
        connection.execute(
            """
            CREATE TABLE IF NOT EXISTS documents (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                knowledge_base_id INTEGER,
                user_id TEXT,
                file_name TEXT,
                file_hash TEXT NOT NULL,
                created_at TEXT NOT NULL,
                FOREIGN KEY (knowledge_base_id) REFERENCES knowledge_bases (id)
            )
            """
        )
        connection.execute(
            """
            CREATE TABLE IF NOT EXISTS knowledge_base (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                knowledge_base_id INTEGER,
                doc_id INTEGER NOT NULL,
                basic_info TEXT NOT NULL,
                education_history TEXT NOT NULL,
                experiences TEXT NOT NULL,
                created_at TEXT NOT NULL,
                updated_at TEXT NOT NULL,
                FOREIGN KEY (doc_id) REFERENCES documents (id),
                FOREIGN KEY (knowledge_base_id) REFERENCES knowledge_bases (id)
            )
            """
        )
        ensure_column(connection, "knowledge_bases", "basic_info", "TEXT")
        ensure_column(connection, "knowledge_bases", "education_history", "TEXT")
        ensure_column(connection, "knowledge_bases", "experiences", "TEXT")
        ensure_column(connection, "documents", "knowledge_base_id", "INTEGER")
        ensure_column(connection, "knowledge_base", "knowledge_base_id", "INTEGER")

        default_base_id = ensure_default_knowledge_base(connection)

        connection.execute(
            "UPDATE knowledge_bases SET basic_info = COALESCE(basic_info, '{}'), education_history = COALESCE(education_history, '[]'), experiences = COALESCE(experiences, '[]')"
        )
        connection.execute(
            "UPDATE knowledge_base SET knowledge_base_id = ? WHERE knowledge_base_id IS NULL",
            (default_base_id,),
        )
        connection.execute(
            "UPDATE documents SET knowledge_base_id = ? WHERE knowledge_base_id IS NULL",
            (default_base_id,),
        )
        connection.commit()
    finally:
        connection.close()



def utc_now_iso() -> str:
    return datetime.utcnow().replace(microsecond=0).isoformat() + "Z"



def ensure_column(connection: sqlite3.Connection, table_name: str, column_name: str, column_sql: str) -> None:
    columns = connection.execute("PRAGMA table_info({0})".format(table_name)).fetchall()
    existing_names = {column["name"] for column in columns}
    if column_name not in existing_names:
        connection.execute(
            "ALTER TABLE {0} ADD COLUMN {1} {2}".format(table_name, column_name, column_sql)
        )



def ensure_default_knowledge_base(connection: sqlite3.Connection) -> int:
    existing = connection.execute(
        "SELECT id FROM knowledge_bases WHERE name = ? LIMIT 1",
        ("?????",),
    ).fetchone()
    if existing is not None:
        return int(existing["id"])

    now = utc_now_iso()
    cursor = connection.execute(
        """
        INSERT INTO knowledge_bases (
            user_id,
            name,
            basic_info,
            education_history,
            experiences,
            created_at,
            updated_at
        ) VALUES (?, ?, ?, ?, ?, ?, ?)
        """,
        (None, "?????", "{}", "[]", "[]", now, now),
    )
    return int(cursor.lastrowid)



def compute_sha256(content: bytes) -> str:
    return hashlib.sha256(content).hexdigest()



def ensure_knowledge_base_exists(knowledge_base_id: int) -> None:
    connection = get_db_connection()
    try:
        row = connection.execute(
            "SELECT id FROM knowledge_bases WHERE id = ?",
            (knowledge_base_id,),
        ).fetchone()
        if row is None:
            raise HTTPException(status_code=404, detail="????????")
    finally:
        connection.close()



def create_knowledge_base(name: str) -> Dict[str, Any]:
    cleaned_name = name.strip()
    if not cleaned_name:
        raise HTTPException(status_code=400, detail="?????????")

    now = utc_now_iso()
    connection = get_db_connection()
    try:
        try:
            cursor = connection.execute(
                """
                INSERT INTO knowledge_bases (
                    user_id,
                    name,
                    basic_info,
                    education_history,
                    experiences,
                    created_at,
                    updated_at
                ) VALUES (?, ?, ?, ?, ?, ?, ?)
                """,
                (None, cleaned_name, "{}", "[]", "[]", now, now),
            )
        except sqlite3.IntegrityError as exc:
            raise HTTPException(status_code=409, detail="????????") from exc
        connection.commit()
        return {
            "id": int(cursor.lastrowid),
            "name": cleaned_name,
            "basic_info": {},
            "education_history": [],
            "experiences": [],
            "created_at": now,
            "updated_at": now,
        }
    finally:
        connection.close()


def rename_knowledge_base(knowledge_base_id: int, name: str) -> Dict[str, Any]:
    cleaned_name = name.strip()
    if not cleaned_name:
        raise HTTPException(status_code=400, detail="知识库名称不能为空")

    now = utc_now_iso()
    connection = get_db_connection()
    try:
        existing = connection.execute(
            "SELECT id FROM knowledge_bases WHERE id = ?",
            (knowledge_base_id,),
        ).fetchone()
        if existing is None:
            raise HTTPException(status_code=404, detail="知识库不存在")

        try:
            connection.execute(
                """
                UPDATE knowledge_bases
                SET name = ?, updated_at = ?
                WHERE id = ?
                """,
                (cleaned_name, now, knowledge_base_id),
            )
        except sqlite3.IntegrityError as exc:
            raise HTTPException(status_code=409, detail="知识库名称已存在") from exc

        connection.commit()
        return {
            "id": int(knowledge_base_id),
            "name": cleaned_name,
            "updated_at": now,
        }
    finally:
        connection.close()



def create_document_record(
    *,
    knowledge_base_id: int,
    user_id: Optional[str],
    file_name: Optional[str],
    file_hash: str,
) -> int:
    connection = get_db_connection()
    try:
        cursor = connection.execute(
            """
            INSERT INTO documents (knowledge_base_id, user_id, file_name, file_hash, created_at)
            VALUES (?, ?, ?, ?, ?)
            """,
            (knowledge_base_id, user_id, file_name, file_hash, utc_now_iso()),
        )
        connection.commit()
        return int(cursor.lastrowid)
    finally:
        connection.close()


def save_knowledge_record(payload: KnowledgeSavePayload) -> Dict[str, Any]:
    ensure_knowledge_base_exists(payload.knowledge_base_id)

    now = utc_now_iso()
    connection = get_db_connection()
    try:
        basic_info_json = json.dumps(payload.basic_info, ensure_ascii=False)
        education_history_json = json.dumps(payload.education_history, ensure_ascii=False)
        experiences_json = json.dumps(payload.experiences, ensure_ascii=False)

        connection.execute(
            """
            UPDATE knowledge_bases
            SET basic_info = ?, education_history = ?, experiences = ?, updated_at = ?
            WHERE id = ?
            """,
            (
                basic_info_json,
                education_history_json,
                experiences_json,
                now,
                payload.knowledge_base_id,
            ),
        )
        connection.commit()
        return {
            "knowledge_base_id": int(payload.knowledge_base_id),
        }
    finally:
        connection.close()



def list_knowledge_bases() -> List[Dict[str, Any]]:
    connection = get_db_connection()
    try:
        rows = connection.execute(
            """
            SELECT
                kb.id,
                kb.user_id,
                kb.name,
                kb.basic_info,
                kb.education_history,
                kb.experiences,
                kb.created_at,
                kb.updated_at,
                COUNT(d.id) AS document_count
            FROM knowledge_bases kb
            LEFT JOIN documents d ON d.knowledge_base_id = kb.id
            GROUP BY kb.id, kb.user_id, kb.name, kb.basic_info, kb.education_history, kb.experiences, kb.created_at, kb.updated_at
            ORDER BY kb.updated_at DESC, kb.id DESC
            """
        ).fetchall()
        return [
            {
                "id": int(row["id"]),
                "user_id": row["user_id"],
                "name": row["name"],
                "created_at": row["created_at"],
                "updated_at": row["updated_at"],
                "document_count": int(row["document_count"] or 0),
                "data": {
                    "basic_info": json.loads(row["basic_info"] or "{}"),
                    "education_history": json.loads(row["education_history"] or "[]"),
                    "experiences": json.loads(row["experiences"] or "[]"),
                },
            }
            for row in rows
        ]
    finally:
        connection.close()



def get_knowledge_base_content(knowledge_base_id: int) -> Dict[str, Any]:
    connection = get_db_connection()
    try:
        row = connection.execute(
            """
            SELECT name, basic_info, education_history, experiences
            FROM knowledge_bases
            WHERE id = ?
            """,
            (knowledge_base_id,),
        ).fetchone()
        if row is None:
            raise HTTPException(status_code=404, detail="????????")

        return {
            "name": row["name"],
            "basic_info": json.loads(row["basic_info"] or "{}"),
            "education_history": json.loads(row["education_history"] or "[]"),
            "experiences": json.loads(row["experiences"] or "[]"),
        }
    finally:
        connection.close()



def delete_knowledge_base(knowledge_base_id: int) -> None:
    connection = get_db_connection()
    try:
        existing = connection.execute(
            "SELECT id FROM knowledge_bases WHERE id = ?",
            (knowledge_base_id,),
        ).fetchone()
        if existing is None:
            raise HTTPException(status_code=404, detail="??????")

        connection.execute(
            "DELETE FROM documents WHERE knowledge_base_id = ?",
            (knowledge_base_id,),
        )
        connection.execute(
            "DELETE FROM knowledge_base WHERE knowledge_base_id = ?",
            (knowledge_base_id,),
        )
        connection.execute(
            "DELETE FROM knowledge_bases WHERE id = ?",
            (knowledge_base_id,),
        )
        connection.commit()
    finally:
        connection.close()


def normalize_text_input(text: str) -> str:
    normalized = text.replace("\r\n", "\n").replace("\r", "\n").strip()
    if not normalized:
        raise HTTPException(status_code=400, detail="输入文本为空")
    return normalized


PLACEHOLDER_PATTERN = re.compile(r"\{\{\s*([^{}]+?)\s*\}\}")


def resolve_template_path(data: Any, path: str) -> Any:
    tokens = re.findall(r"([^.\[\]]+)|\[(\d+)\]", path)
    current = data
    for name_token, index_token in tokens:
        token = name_token if name_token else int(index_token)
        if isinstance(token, int):
            if not isinstance(current, list) or token >= len(current):
                return None
            current = current[token]
        else:
            if not isinstance(current, dict):
                return None
            current = current.get(token)
        if current is None:
            return None
    return current


def stringify_template_value(value: Any) -> str:
    if value is None:
        return ""
    if isinstance(value, (str, int, float)):
        return str(value)
    if isinstance(value, list):
        if not value:
            return ""
        if all(not isinstance(item, (dict, list)) for item in value):
            return "?".join(str(item) for item in value if item not in (None, ""))
        return json.dumps(value, ensure_ascii=False)
    if isinstance(value, dict):
        return json.dumps(value, ensure_ascii=False)
    return str(value)


def flatten_template_context(data: Any, prefix: str = "") -> Dict[str, str]:
    flattened: Dict[str, str] = {}

    if isinstance(data, dict):
        for key, value in data.items():
            next_prefix = "{0}.{1}".format(prefix, key) if prefix else str(key)
            flattened.update(flatten_template_context(value, next_prefix))
    elif isinstance(data, list):
        flattened[prefix] = stringify_template_value(data)
        for index, value in enumerate(data):
            next_prefix = "{0}[{1}]".format(prefix, index) if prefix else "[{0}]".format(index)
            flattened.update(flatten_template_context(value, next_prefix))
    else:
        flattened[prefix] = stringify_template_value(data)

    return flattened


def render_template_text(text: str, context: Dict[str, Any]) -> str:
    def replacer(match):
        raw_path = match.group(1).strip()
        value = resolve_template_path(context, raw_path)
        return stringify_template_value(value)

    return PLACEHOLDER_PATTERN.sub(replacer, text)


def replace_in_paragraph(paragraph: Any, context: Dict[str, Any]) -> None:
    if not paragraph.text:
        return
    rendered = render_template_text(paragraph.text, context)
    if rendered == paragraph.text:
        return
    if paragraph.runs:
        paragraph.runs[0].text = rendered
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(rendered)


def replace_in_table(table: Any, context: Dict[str, Any]) -> None:
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                replace_in_paragraph(paragraph, context)
            for nested_table in cell.tables:
                replace_in_table(nested_table, context)


def set_cell_text(cell: Any, text: str) -> None:
    if cell is None:
        return
    cell.text = text or ""


def select_experiences_by_categories(experiences: List[Dict[str, Any]], categories: Tuple[str, ...]) -> List[Dict[str, Any]]:
    return [item for item in experiences if str(item.get("category") or "").strip() in categories]


def build_application_form_payload(context: Dict[str, Any]) -> Dict[str, Any]:
    basic_info = context.get("basic_info") or {}
    education_history = context.get("education_history") or []
    experiences = context.get("experiences") or []
    primary_education = education_history[0] if education_history else {}

    awards = select_experiences_by_categories(experiences, ("荣誉奖项", "比赛经历"))
    projects = select_experiences_by_categories(experiences, ("科研项目", "实习经历", "工作经历", "社会实践", "学生工作"))

    grade_parts = [basic_info.get("current_status") or "", basic_info.get("major") or ""]
    ranking_parts = [primary_education.get("ranking") or "", primary_education.get("gpa") or ""]

    project_lines: List[str] = []
    for item in projects[:5]:
        line = "；".join(
            [
                str(item.get("name") or "").strip(),
                str(item.get("achievement") or "").strip(),
                str(item.get("description") or "").strip(),
            ]
        ).strip("；")
        if line:
            project_lines.append(line)

    deed_lines: List[str] = []
    for item in experiences[:8]:
        parts = [
            str(item.get("name") or "").strip(),
            str(item.get("role_or_title") or "").strip(),
            str(item.get("achievement") or "").strip(),
            str(item.get("description") or "").strip(),
        ]
        line = "，".join([part for part in parts if part]).strip("，")
        if line:
            deed_lines.append(line)

    return {
        "姓名": str(basic_info.get("name") or ""),
        "性别": str(basic_info.get("gender") or ""),
        "出生年月": str(basic_info.get("birth_date") or ""),
        "在读学历": str(basic_info.get("degree") or ""),
        "学院/研究院": str(basic_info.get("university") or ""),
        "年级专业": " ".join([part for part in grade_parts if part]).strip(),
        "手机号码": str(basic_info.get("phone") or ""),
        "电子邮箱": str(basic_info.get("email") or ""),
        "上一学年综合测评/各维度评价结果": " / ".join([part for part in ranking_parts if part]).strip(),
        "科研成果": "\n".join(project_lines),
        "个人事迹": "\n".join(deed_lines),
        "获奖情况": awards[:10],
    }


def fill_known_application_form_tables(document: Any, context: Dict[str, Any]) -> None:
    payload = build_application_form_payload(context)
    awards = payload["获奖情况"]

    for table in document.tables:
        for row_index, row in enumerate(table.rows):
            cells = row.cells
            texts = [cell.text.replace("\n", " ").strip() for cell in cells]
            joined = " | ".join(texts)

            if row_index == 0 and "姓名" in texts:
                set_cell_text(cells[1], payload["姓名"])
                set_cell_text(cells[4], payload["性别"])
                set_cell_text(cells[6], payload["出生年月"])
            elif row_index == 1 and "在读学历" in texts:
                set_cell_text(cells[1], "")
                set_cell_text(cells[4], payload["在读学历"])
                set_cell_text(cells[6], "")
            elif row_index == 2 and "现任职务" in texts:
                set_cell_text(cells[1], "")
                set_cell_text(cells[5], "")
            elif row_index == 3 and "学院/研究院" in joined:
                set_cell_text(cells[1], payload["学院/研究院"])
                set_cell_text(cells[6], payload["年级专业"])
            elif row_index == 4 and "手机号码" in joined:
                set_cell_text(cells[1], payload["手机号码"])
                set_cell_text(cells[6], payload["电子邮箱"])
            elif row_index == 5 and "上一学年综合测评" in joined:
                set_cell_text(cells[1], "")
                set_cell_text(cells[6], payload["上一学年综合测评/各维度评价结果"])
            elif row_index >= 8 and row_index <= 17 and "入学以来 获奖情况" in joined:
                award_index = row_index - 8
                award = awards[award_index] if award_index < len(awards) else None
                if award:
                    set_cell_text(cells[1], str(award_index + 1))
                    set_cell_text(cells[2], str(award.get("end_date") or award.get("start_date") or ""))
                    set_cell_text(cells[4], str(award.get("name") or award.get("achievement") or ""))
                    set_cell_text(cells[7], str(award.get("organization") or ""))
                else:
                    for idx in range(1, len(cells)):
                        set_cell_text(cells[idx], "")
            elif row_index == 18 and "科研成果" in joined:
                set_cell_text(cells[1], payload["科研成果"])
            elif row_index == 19 and "个人事迹" in joined:
                set_cell_text(cells[1], payload["个人事迹"])


def apply_xml_placeholder_replacement(docx_bytes: bytes, context: Dict[str, Any]) -> bytes:
    flat_context = flatten_template_context(context)
    if not flat_context:
        return docx_bytes

    source_buffer = io.BytesIO(docx_bytes)
    output_buffer = io.BytesIO()

    with zipfile.ZipFile(source_buffer, "r") as source_zip, zipfile.ZipFile(output_buffer, "w", zipfile.ZIP_DEFLATED) as target_zip:
        for item in source_zip.infolist():
            data = source_zip.read(item.filename)
            if item.filename.startswith("word/") and item.filename.endswith(".xml"):
                try:
                    xml_text = data.decode("utf-8")
                except UnicodeDecodeError:
                    target_zip.writestr(item, data)
                    continue

                for key, value in flat_context.items():
                    xml_text = xml_text.replace("{{{{{0}}}}}".format(key), value)
                    xml_text = xml_text.replace("{{{{ {0} }}}}".format(key), value)

                target_zip.writestr(item, xml_text.encode("utf-8"))
            else:
                target_zip.writestr(item, data)

    output_buffer.seek(0)
    return output_buffer.getvalue()


def convert_word_via_com(file_bytes: bytes, original_suffix: str) -> bytes:
    try:
        import pythoncom
        from win32com.client import DispatchEx
    except Exception as exc:
        raise HTTPException(
            status_code=500,
            detail="当前环境不支持 Word 模板转换，请安装并启用 Microsoft Word 自动化组件: {0}".format(exc),
        ) from exc

    with tempfile.TemporaryDirectory() as temp_dir:
        source_path = os.path.join(temp_dir, "source{0}".format(original_suffix))
        target_path = os.path.join(temp_dir, "converted.docx")

        with open(source_path, "wb") as file_obj:
            file_obj.write(file_bytes)

        pythoncom.CoInitialize()
        word = None
        document = None
        try:
            word = DispatchEx("Word.Application")
            word.Visible = False
            word.DisplayAlerts = 0
            document = word.Documents.Open(source_path, ReadOnly=True)
            document.SaveAs(target_path, FileFormat=16)
        except Exception as exc:
            raise HTTPException(
                status_code=400,
                detail="Word 模板转换失败，请确认本机已安装 Microsoft Word，且文件格式有效: {0}".format(exc),
            ) from exc
        finally:
            if document is not None:
                try:
                    document.Close(False)
                except Exception:
                    pass
            if word is not None:
                try:
                    word.Quit()
                except Exception:
                    pass
            pythoncom.CoUninitialize()

        with open(target_path, "rb") as file_obj:
            return file_obj.read()


def prepare_word_template_bytes(file_bytes: bytes, filename: str) -> bytes:
    lower_name = filename.lower()
    is_zip_docx = zipfile.is_zipfile(io.BytesIO(file_bytes))

    if lower_name.endswith(".docx") and is_zip_docx:
        return file_bytes

    if lower_name.endswith(".doc") or lower_name.endswith(".docx"):
        return convert_word_via_com(file_bytes, ".doc" if lower_name.endswith(".doc") else ".docx")

    raise HTTPException(status_code=400, detail="仅支持上传 DOC 或 DOCX 模板文件")


def fill_word_template(file_bytes: bytes, knowledge_base_id: int, original_filename: str) -> Tuple[bytes, str]:
    normalized_bytes = prepare_word_template_bytes(file_bytes, original_filename)
    try:
        document = Document(io.BytesIO(normalized_bytes))
    except Exception as exc:
        raise HTTPException(status_code=400, detail="Word 模板解析失败: {0}".format(exc)) from exc

    context = get_knowledge_base_content(knowledge_base_id)

    for paragraph in document.paragraphs:
        replace_in_paragraph(paragraph, context)

    for table in document.tables:
        replace_in_table(table, context)

    fill_known_application_form_tables(document, context)

    for section in document.sections:
        for paragraph in section.header.paragraphs:
            replace_in_paragraph(paragraph, context)
        for table in section.header.tables:
            replace_in_table(table, context)
        for paragraph in section.footer.paragraphs:
            replace_in_paragraph(paragraph, context)
        for table in section.footer.tables:
            replace_in_table(table, context)

    output = io.BytesIO()
    document.save(output)
    output.seek(0)
    final_bytes = apply_xml_placeholder_replacement(output.getvalue(), context)
    template_name = context.get("name") or "knowledge_base"
    safe_name = re.sub(r"[^\w\u4e00-\u9fff-]+", "_", template_name).strip("_") or "knowledge_base"
    return final_bytes, f"{safe_name}_filled.docx"


def build_download_headers(filename: str) -> Dict[str, str]:
    ascii_name = re.sub(r"[^A-Za-z0-9._-]+", "_", filename).strip("_") or "filled_template.docx"
    encoded_name = urllib.parse.quote(filename)
    return {
        "Content-Disposition": "attachment; filename=\"{0}\"; filename*=UTF-8''{1}".format(
            ascii_name,
            encoded_name,
        )
    }


def clean_extracted_text(text: str) -> str:
    lines = [line.strip() for line in text.replace("\r\n", "\n").replace("\r", "\n").split("\n")]
    non_empty_lines = [line for line in lines if line]
    if not non_empty_lines:
        raise HTTPException(status_code=400, detail="未能提取到有效文本内容")

    merged = "\n".join(non_empty_lines)
    if len(merged) > MAX_TEXT_CHARS:
        merged = merged[:MAX_TEXT_CHARS]
    return merged


def decode_text_file(file_bytes: bytes) -> str:
    for encoding in ("utf-8", "utf-8-sig", "gb18030", "gbk"):
        try:
            return clean_extracted_text(file_bytes.decode(encoding))
        except UnicodeDecodeError:
            continue
    raise HTTPException(status_code=400, detail="TXT 文件编码无法识别，请使用 UTF-8 或 GBK")


def extract_docx_text(file_bytes: bytes) -> str:
    try:
        document = Document(io.BytesIO(file_bytes))
    except Exception as exc:
        raise HTTPException(status_code=400, detail="Word 文件解析失败: {0}".format(exc)) from exc

    paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
    rows: List[str] = []
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                rows.append(" | ".join(cells))

    content = "\n".join(paragraphs + rows).strip()
    if not content:
        raise HTTPException(status_code=400, detail="Word 文件内容为空")
    return clean_extracted_text(content)


def pil_image_from_bytes(file_bytes: bytes) -> Image.Image:
    try:
        return Image.open(io.BytesIO(file_bytes)).convert("RGB")
    except Exception as exc:
        raise HTTPException(status_code=400, detail="图片解析失败: {0}".format(exc)) from exc


def image_to_data_url(image: Image.Image) -> str:
    buffer = io.BytesIO()
    image.save(buffer, format="JPEG", quality=92)
    encoded = buffer.getvalue()
    return "data:image/jpeg;base64,{0}".format(base64.b64encode(encoded).decode("utf-8"))


def pdf_first_page_to_image(file_bytes: bytes) -> Image.Image:
    try:
        pdf = fitz.open(stream=file_bytes, filetype="pdf")
    except Exception as exc:
        raise HTTPException(status_code=400, detail="PDF 解析失败: {0}".format(exc)) from exc

    try:
        if pdf.page_count < 1:
            raise HTTPException(status_code=400, detail="PDF 为空，无法提取首页")
        page = pdf.load_page(0)
        pix = page.get_pixmap(matrix=fitz.Matrix(1.8, 1.8), alpha=False)
        return Image.open(io.BytesIO(pix.tobytes("png"))).convert("RGB")
    finally:
        pdf.close()


def get_ocr_engine():
    global _ocr_engine
    if _ocr_engine is not None:
        return _ocr_engine

    try:
        from paddleocr import PaddleOCR
    except Exception as exc:
        raise HTTPException(
            status_code=503,
            detail=(
                "本地 OCR 暂不可用，请先安装并配置 PaddleOCR / PaddlePaddle，"
                "或暂时使用 TXT、DOCX、直接文本输入。原始错误: {0}"
            ).format(exc),
        ) from exc

    try:
        _ocr_engine = PaddleOCR(lang=OCR_LANG)
    except Exception as exc:
        raise HTTPException(
            status_code=503,
            detail="OCR 引擎初始化失败，请检查 PaddlePaddle 环境。原始错误: {0}".format(exc),
        ) from exc

    return _ocr_engine


def run_ocr_on_image(image: Image.Image) -> str:
    ocr = get_ocr_engine()
    image_array = np.array(image)

    try:
        result = ocr.ocr(image_array)
    except Exception as exc:
        raise HTTPException(status_code=500, detail="OCR 识别失败: {0}".format(exc)) from exc

    texts: List[str] = []
    for page in result or []:
        for line in page or []:
            if len(line) < 2:
                continue
            text = line[1][0].strip()
            if text:
                texts.append(text)

    if not texts:
        raise HTTPException(status_code=400, detail="OCR 未识别到有效文本")
    return clean_extracted_text("\n".join(texts))


def extract_pdf_text(file_bytes: bytes) -> str:
    try:
        pdf = fitz.open(stream=file_bytes, filetype="pdf")
    except Exception as exc:
        raise HTTPException(status_code=400, detail="PDF 解析失败: {0}".format(exc)) from exc

    try:
        page_texts: List[str] = []
        for page in pdf:
            text = page.get_text("text").strip()
            if text:
                page_texts.append(text)

        merged = "\n".join(page_texts).strip()
        if len(merged) >= PDF_TEXT_THRESHOLD:
            return clean_extracted_text(merged)

        ocr_texts: List[str] = []
        for index in range(pdf.page_count):
            page = pdf.load_page(index)
            pix = page.get_pixmap(matrix=fitz.Matrix(1.8, 1.8), alpha=False)
            image = Image.open(io.BytesIO(pix.tobytes("png"))).convert("RGB")
            ocr_texts.append(run_ocr_on_image(image))
        return clean_extracted_text("\n".join(ocr_texts))
    finally:
        pdf.close()


def build_schema_description() -> Dict[str, Any]:
    return {
        "basic_info": {
            "name": "string or null",
            "gender": "string or null",
            "phone": "string or null",
            "email": "string or null",
            "birth_date": "YYYY-MM-DD or null",
            "university": "string or null",
            "major": "string or null",
            "degree": "string or null",
            "graduation_year": "string or null",
            "current_status": "string or null",
        },
        "education_history": [
            {
                "institution": "学校名称",
                "major": "专业",
                "degree": "学位（本科/硕士/博士）",
                "start_date": "YYYY-MM or null",
                "end_date": "YYYY-MM or null",
                "gpa": "string or null",
                "ranking": "排名（如前10%） or null",
                "courses": ["课程名1", "课程名2"],
            }
        ],
        "experiences": [
            {
                "category": "枚举值：[实习经历, 工作经历, 比赛经历, 科研项目, 荣誉奖项, 学生工作, 社会实践, 技能证书, 语言证书, 其他]",
                "name": "名称（比赛/公司/证书名）",
                "role_or_title": "角色或职位（队长/实习生/一等奖）",
                "organization": "组织单位（学校/公司/主办方）",
                "level": "级别（国家级/省级等，非赛事类可null）",
                "start_date": "YYYY-MM or null",
                "end_date": "YYYY-MM or null",
                "achievement": "具体成果（排名、分数、奖金等）",
                "description": "详细描述",
            }
        ],
    }


def build_messages(plain_text: str) -> List[Dict[str, Any]]:
    schema = json.dumps(build_schema_description(), ensure_ascii=False)
    system_prompt = (
        "你是一个极其精准的个人履历数据提取专家。"
        "用户会输入一段由文本、Word、PDF 或图片 OCR 提取出的文本材料，例如比赛证书、成绩单、项目总结报告、学生工作材料等。"
        "你需要分析内容，并严格按照指定 JSON 结构输出提取到的信息。"
        "如果某项信息不存在，请将其值设为 null，不要自行编造。"
        "你必须只输出 JSON 对象，不能输出 Markdown、解释、代码块或其他文本。"
        "basic_info 必须始终存在，字段缺失时填 null。"
        "education_history 必须始终为数组；如果没有识别到教育经历，返回空数组。"
        "experiences 必须始终为数组；如果无法识别出任何经历，返回空数组。"
        "experience.category 只能从以下枚举中选择一个：实习经历、工作经历、比赛经历、科研项目、荣誉奖项、学生工作、社会实践、技能证书、语言证书、其他。"
        "birth_date 必须使用 YYYY-MM-DD 格式，无法确定则填 null。"
        "education_history 中的 start_date 和 end_date 必须使用 YYYY-MM 格式，无法确定则填 null。"
        "experiences 中的 start_date 和 end_date 必须使用 YYYY-MM 格式，无法确定则填 null。"
        "如果单个材料中包含多段经历，请拆成 experiences 中的多条记录。"
        "education_history.courses 必须始终为数组，没有则返回空数组。"
        "experiences.description 应基于原文做精炼总结，控制在 1 到 3 句内，优先保留职责、方法和结果，不要写成超长段落。"
        "当你提取技能证书类信息时，如果原文中多个技能用逗号、顿号、分号、'、' 或类似分隔符连接，请拆分成多个独立的 experience 条目。"
        "拆分后的每个技能条目 category 填 '技能证书'，name 只保留单个技能名称，description 或 achievement 填熟练程度，例如'熟练使用'。"
        "当你提取语言证书类信息时，例如 CET-4、CET-6、雅思、托福，每个证书必须单独生成一条 experience。"
        "语言证书条目的 category 填 '语言证书'，name 填证书名称，achievement 填分数或等级。"
        "不要额外增加未定义字段。"
        "输出字段结构必须为: {0}"
    ).format(schema)

    user_prompt = (
        "请分析以下材料，并严格输出指定 JSON 结构。"
        "如果某项信息不存在，必须填 null，不要编造。"
        "请尽可能抽取 basic_info、education_history，以及一条或多条 experiences。"
        "如果材料中包含教育背景信息，请放入 education_history。"
        "如果材料只对应一段经历，就输出一个 experiences 元素。"
        "如果出现技能列表，请拆成多个 category='技能证书' 的条目。"
        "如果出现语言考试或语言证书，请按证书逐条拆分，并把分数写入 achievement。"
        "直接输出 JSON 字符串，不要包含任何 Markdown 语法。\n\n"
        "材料内容如下：\n{0}"
    ).format(plain_text)

    return [
        {"role": "system", "content": system_prompt},
        {"role": "user", "content": user_prompt},
    ]


def build_vision_messages(image_data_url: str) -> List[Dict[str, Any]]:
    schema = json.dumps(build_schema_description(), ensure_ascii=False)
    system_prompt = (
        "你是一个极其精准的个人履历数据提取专家。"
        "用户会输入一张证明材料图片或 PDF 首页截图。"
        "你需要分析内容，并严格按照指定 JSON 结构输出提取到的信息。"
        "如果某项信息不存在，请将其值设为 null，不要自行编造。"
        "你必须只输出 JSON 对象，不能输出 Markdown、解释、代码块或其他文本。"
        "basic_info 必须始终存在，字段缺失时填 null。"
        "education_history 必须始终为数组；如果没有识别到教育经历，返回空数组。"
        "experiences 必须始终为数组；如果无法识别出任何经历，返回空数组。"
        "experience.category 只能从以下枚举中选择一个：实习经历、工作经历、比赛经历、科研项目、荣誉奖项、学生工作、社会实践、技能证书、语言证书、其他。"
        "birth_date 必须使用 YYYY-MM-DD 格式，无法确定则填 null。"
        "education_history 中的 start_date 和 end_date 必须使用 YYYY-MM 格式，无法确定则填 null。"
        "experiences 中的 start_date 和 end_date 必须使用 YYYY-MM 格式，无法确定则填 null。"
        "如果单个材料中包含多段经历，请拆成 experiences 中的多条记录。"
        "education_history.courses 必须始终为数组，没有则返回空数组。"
        "experiences.description 应基于原文做精炼总结，控制在 1 到 3 句内，优先保留职责、方法和结果，不要写成超长段落。"
        "当你提取技能证书类信息时，如果原文中多个技能用逗号、顿号、分号、'、' 或类似分隔符连接，请拆分成多个独立的 experience 条目。"
        "拆分后的每个技能条目 category 填 '技能证书'，name 只保留单个技能名称，description 或 achievement 填熟练程度，例如'熟练使用'。"
        "当你提取语言证书类信息时，例如 CET-4、CET-6、雅思、托福，每个证书必须单独生成一条 experience。"
        "语言证书条目的 category 填 '语言证书'，name 填证书名称，achievement 填分数或等级。"
        "不要额外增加未定义字段。"
        "输出字段结构必须为: {0}"
    ).format(schema)

    return [
        {"role": "system", "content": system_prompt},
        {
            "role": "user",
            "content": [
                {
                    "type": "text",
                    "text": (
                        "请分析这份材料，并严格输出指定 JSON 结构。"
                        "如果某项信息不存在，必须填 null，不要编造。"
                        "请尽可能抽取 basic_info、education_history，以及一条或多条 experiences。"
                        "如果材料中包含教育背景信息，请放入 education_history。"
                        "如果材料只对应一段经历，就输出一个 experiences 元素。"
                        "如果出现技能列表，请拆成多个 category='技能证书' 的条目。"
                        "如果出现语言考试或语言证书，请按证书逐条拆分，并把分数写入 achievement。"
                        "直接输出 JSON 字符串，不要包含任何 Markdown 语法。"
                    ),
                },
                {
                    "type": "image_url",
                    "image_url": {"url": image_data_url},
                },
            ],
        },
    ]


def create_ark_client() -> Ark:
    if not ARK_API_KEY:
        raise HTTPException(status_code=500, detail="缺少环境变量 ARK_API_KEY")
    if not ARK_TEXT_ENDPOINT_ID:
        raise HTTPException(
            status_code=500,
            detail="缺少文本模型配置。请在 .env 中设置 ARK_TEXT_ENDPOINT_ID，或使用 DOUBAO_TEXT_MODEL / ARK_ENDPOINT_ID。",
        )

    http_client = httpx.Client(timeout=ARK_TIMEOUT, trust_env=False)
    return Ark(
        base_url=ARK_BASE_URL,
        api_key=ARK_API_KEY,
        timeout=ARK_TIMEOUT,
        http_client=http_client,
    )


def extract_json_candidate(text: str) -> str:
    stripped = text.strip()
    if stripped.startswith("```"):
        fenced = re.search(r"```(?:json)?\s*(.*?)\s*```", stripped, re.DOTALL | re.IGNORECASE)
        if fenced:
            stripped = fenced.group(1).strip()

    start = stripped.find("{")
    if start == -1:
        return stripped

    depth = 0
    in_string = False
    escape = False
    for index, char in enumerate(stripped[start:], start=start):
        if in_string:
            if escape:
                escape = False
            elif char == "\\":
                escape = True
            elif char == '"':
                in_string = False
        else:
            if char == '"':
                in_string = True
            elif char == "{":
                depth += 1
            elif char == "}":
                depth -= 1
                if depth == 0:
                    return stripped[start : index + 1]
    return stripped[start:]


def parse_model_json(content: str) -> Dict[str, Any]:
    candidates = [content.strip(), extract_json_candidate(content)]
    for candidate in candidates:
        if not candidate:
            continue
        try:
            return json.loads(candidate)
        except json.JSONDecodeError:
            continue
    logger.error("Model returned non-JSON content or truncated JSON: %s", content)
    raise HTTPException(status_code=502, detail="大模型未返回合法 JSON")


def split_text_items(value: str) -> List[str]:
    parts = re.split(r"[、,，;；]+", value)
    return [part.strip(" \t\r\n\"'“”") for part in parts if part.strip(" \t\r\n\"'“”")]


def normalize_skill_and_language_experiences(data: Dict[str, Any]) -> Dict[str, Any]:
    experiences = data.get("experiences")
    if not isinstance(experiences, list):
        return data

    normalized: List[Dict[str, Any]] = []
    language_keywords = ("CET-4", "CET-6", "雅思", "托福", "IELTS", "TOEFL", "TEM-4", "TEM-8")

    for item in experiences:
        if not isinstance(item, dict):
            continue

        category = str(item.get("category") or "").strip()
        name = str(item.get("name") or "").strip()
        achievement = item.get("achievement")
        description = item.get("description")

        if category == "技能证书" and name:
            items = split_text_items(name)
            if len(items) > 1:
                base_detail = description or achievement or None
                for skill_name in items:
                    new_item = dict(item)
                    new_item["name"] = skill_name
                    if base_detail and not new_item.get("description"):
                        new_item["description"] = base_detail
                    normalized.append(new_item)
                continue

        if category == "语言证书" and name:
            items = split_text_items(name)
            if len(items) > 1:
                for cert_name in items:
                    new_item = dict(item)
                    new_item["name"] = cert_name
                    normalized.append(new_item)
                continue

        if category != "语言证书" and name and any(keyword in name for keyword in language_keywords):
            items = split_text_items(name)
            if len(items) > 1:
                for cert_name in items:
                    cert_item = dict(item)
                    cert_item["category"] = "语言证书"
                    cert_item["name"] = cert_name
                    normalized.append(cert_item)
                continue

        normalized.append(item)

    data["experiences"] = normalized
    return data


def call_text_model(plain_text: str) -> Dict[str, Any]:
    client = create_ark_client()
    messages = build_messages(plain_text)

    try:
        completion = client.chat.completions.create(
            model=ARK_TEXT_ENDPOINT_ID,
            messages=messages,
            temperature=0,
            max_tokens=ARK_MAX_TOKENS,
            response_format={"type": "json_object"},
        )
    except Exception as exc:
        logger.exception("Ark text model call failed")
        message = str(exc)
        lower_message = message.lower()
        if "timeout" in lower_message or "timed out" in lower_message:
            raise HTTPException(
                status_code=504,
                detail="大模型调用超时，请稍后重试，或调大 ARK_TIMEOUT / 减小输入内容长度",
            ) from exc
        raise HTTPException(status_code=502, detail="大模型调用失败: {0}".format(message)) from exc

    try:
        content = completion.choices[0].message.content
    except Exception as exc:
        logger.exception("Unexpected Ark response structure: %r", completion)
        raise HTTPException(status_code=502, detail="大模型返回结果结构异常") from exc

    if not content:
        raise HTTPException(status_code=502, detail="大模型未返回有效内容")

    return normalize_skill_and_language_experiences(parse_model_json(content))


def call_vision_model(image: Image.Image) -> Dict[str, Any]:
    if not ARK_VISION_ENDPOINT_ID:
        raise HTTPException(
            status_code=503,
            detail="OCR 不可用，且未配置视觉模型 ARK_VISION_ENDPOINT_ID / DOUBAO_VISION_MODEL，无法解析图片。",
        )

    client = create_ark_client()
    messages = build_vision_messages(image_to_data_url(image))

    try:
        completion = client.chat.completions.create(
            model=ARK_VISION_ENDPOINT_ID,
            messages=messages,
            temperature=0,
            max_tokens=ARK_MAX_TOKENS,
            response_format={"type": "json_object"},
        )
    except Exception as exc:
        logger.exception("Ark vision model call failed")
        message = str(exc)
        lower_message = message.lower()
        if "timeout" in lower_message or "timed out" in lower_message:
            raise HTTPException(status_code=504, detail="视觉模型调用超时，请稍后重试。") from exc
        raise HTTPException(status_code=502, detail="视觉模型调用失败: {0}".format(message)) from exc

    try:
        content = completion.choices[0].message.content
    except Exception as exc:
        raise HTTPException(status_code=502, detail="视觉模型返回结果结构异常") from exc

    if not content:
        raise HTTPException(status_code=502, detail="视觉模型未返回有效内容")

    return normalize_skill_and_language_experiences(parse_model_json(content))


def detect_input_type(file: Optional[UploadFile], text: Optional[str]) -> str:
    has_file = file is not None
    has_text = bool(text and text.strip())

    if has_file and has_text:
        raise HTTPException(status_code=400, detail="请只提供 file 或 text 其中一种输入")
    if not has_file and not has_text:
        raise HTTPException(status_code=400, detail="请上传文件或直接传入 text 文本")
    return "text" if has_text else "file"


def extract_text_from_file(file: UploadFile, file_bytes: bytes) -> Tuple[Optional[str], str, Optional[Image.Image]]:
    content_type = (file.content_type or "").lower()
    filename = (file.filename or "").lower()

    if content_type in SUPPORTED_TEXT_TYPES or filename.endswith(".txt"):
        return decode_text_file(file_bytes), "plain_text", None

    if content_type in SUPPORTED_DOCX_TYPES or filename.endswith(".docx"):
        return extract_docx_text(file_bytes), "docx_text", None

    if content_type in SUPPORTED_PDF_TYPES or filename.endswith(".pdf"):
        try:
            return extract_pdf_text(file_bytes), "pdf_text_or_ocr", None
        except HTTPException as exc:
            if exc.status_code == 503:
                return None, "pdf_vision_fallback", pdf_first_page_to_image(file_bytes)
            raise

    if content_type in SUPPORTED_IMAGE_TYPES or filename.endswith((".png", ".jpg", ".jpeg")):
        image = pil_image_from_bytes(file_bytes)
        try:
            return run_ocr_on_image(image), "image_ocr", None
        except HTTPException as exc:
            if exc.status_code == 503:
                return None, "image_vision_fallback", image
            raise

    raise HTTPException(
        status_code=400,
        detail="仅支持 PNG、JPG/JPEG、PDF、DOCX、TXT 文件，或直接传入 text 文本",
    )


@app.post("/api/extract")
async def extract_experience(
    file: Optional[UploadFile] = File(None),
    text: Optional[str] = Form(None),
    knowledge_base_id: int = Form(...),
) -> JSONResponse:
    ensure_knowledge_base_exists(knowledge_base_id)
    input_type = detect_input_type(file, text)

    if input_type == "text":
        normalized_text = normalize_text_input(text or "")
        text_hash = compute_sha256(normalized_text.encode("utf-8"))
        doc_id = create_document_record(
            knowledge_base_id=knowledge_base_id,
            user_id=None,
            file_name="direct_input.txt",
            file_hash=text_hash,
        )
        extracted = call_text_model(normalized_text)
        return JSONResponse(
            content={
                "success": True,
                "input_type": "text",
                "text_strategy": "direct_text",
                "doc_id": doc_id,
                "knowledge_base_id": knowledge_base_id,
                "file_hash": text_hash,
                "file_name": "direct_input.txt",
                "data": extracted,
            }
        )

    assert file is not None
    file_bytes = await file.read()
    if not file_bytes:
        raise HTTPException(status_code=400, detail="上传文件为空")

    file_hash = compute_sha256(file_bytes)
    doc_id = create_document_record(
        knowledge_base_id=knowledge_base_id,
        user_id=None,
        file_name=file.filename,
        file_hash=file_hash,
    )

    extracted_text, text_strategy, fallback_image = extract_text_from_file(file, file_bytes)
    if extracted_text is not None:
        extracted = call_text_model(extracted_text)
        extracted_text_preview = extracted_text[:500]
    else:
        assert fallback_image is not None
        extracted = call_vision_model(fallback_image)
        extracted_text_preview = None

    return JSONResponse(
        content={
            "success": True,
            "input_type": "file",
            "doc_id": doc_id,
            "knowledge_base_id": knowledge_base_id,
            "file_hash": file_hash,
            "filename": file.filename,
            "file_name": file.filename,
            "content_type": file.content_type,
            "text_strategy": text_strategy,
            "extracted_text_preview": extracted_text_preview,
            "data": extracted,
        }
    )


@app.post("/api/knowledge/save")
async def save_knowledge(payload: KnowledgeSavePayload) -> JSONResponse:
    saved = save_knowledge_record(payload)
    return JSONResponse(
        content={
            "success": True,
            "message": "保存成功",
            **saved,
        }
    )


@app.get("/api/knowledge-bases")
async def get_knowledge_bases() -> JSONResponse:
    return JSONResponse(
        content={
            "success": True,
            "items": list_knowledge_bases(),
        }
    )


@app.post("/api/knowledge-bases")
async def create_knowledge_base_api(payload: KnowledgeBaseCreatePayload) -> JSONResponse:
    created = create_knowledge_base(payload.name)
    return JSONResponse(
        content={
            "success": True,
            "message": "知识库创建成功",
            "item": created,
        }
    )


@app.put("/api/knowledge-bases/{knowledge_base_id}")
async def rename_knowledge_base_api(
    knowledge_base_id: int, payload: KnowledgeBaseRenamePayload
) -> JSONResponse:
    updated = rename_knowledge_base(knowledge_base_id, payload.name)
    return JSONResponse(
        content={
            "success": True,
            "message": "知识库名称已更新",
            "item": updated,
        }
    )


@app.delete("/api/knowledge-bases/{knowledge_base_id}")
async def delete_knowledge_base_api(knowledge_base_id: int) -> JSONResponse:
    delete_knowledge_base(knowledge_base_id)
    return JSONResponse(
        content={
            "success": True,
            "message": "??????",
        }
    )


@app.post("/api/templates/fill")
async def fill_template(
    file: UploadFile = File(...),
    knowledge_base_id: int = Form(...),
) -> StreamingResponse:
    ensure_knowledge_base_exists(knowledge_base_id)
    filename = (file.filename or "").lower()
    content_type = (file.content_type or "").lower()
    if not (
        filename.endswith(".docx")
        or filename.endswith(".doc")
        or content_type in SUPPORTED_DOCX_TYPES
        or content_type == "application/msword"
    ):
        raise HTTPException(status_code=400, detail="仅支持上传 DOC 或 DOCX 模板文件")

    file_bytes = await file.read()
    if not file_bytes:
        raise HTTPException(status_code=400, detail="上传模板为空")

    filled_bytes, download_name = fill_word_template(file_bytes, knowledge_base_id, file.filename or "template.docx")
    headers = build_download_headers(download_name)
    return StreamingResponse(
        io.BytesIO(filled_bytes),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers=headers,
    )


@app.get("/health")
async def health() -> Dict[str, Any]:
    return {"status": "ok"}


if __name__ == "__main__":
    import uvicorn

    uvicorn.run("main:app", host="0.0.0.0", port=8000, reload=True)
